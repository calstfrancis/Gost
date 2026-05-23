"""wordgen.py — Generate a styled DOCX placeholder template from a Gost state dict."""

from __future__ import annotations
import io
from typing import Dict

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def available() -> bool:
    return _DOCX_AVAILABLE


_FONT_MAP: Dict[str, str] = {
    "times":      "Times New Roman",
    "ebgaramond": "EB Garamond",
    "lmodern":    "Times New Roman",
    "libertine":  "Linux Libertine",
    "palatino":   "Palatino Linotype",
    "garamond":   "Garamond",
    "junicode":   "Junicode",
    "none":       "Times New Roman",
}

# string keys resolved at runtime to avoid referencing WD_ enums at import time
_HEAD_DEFS: Dict[str, Dict] = {
    "SBL":      {"h1_align": "center", "h1_bold": True,  "h1_italic": True,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": False},
    "Chicago":  {"h1_align": "center", "h1_bold": True,  "h1_italic": False,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": False},
    "MLA":      {"h1_align": "left",   "h1_bold": True,  "h1_italic": False,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": False},
    "APA":      {"h1_align": "center", "h1_bold": True,  "h1_italic": False,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": True},
    "ASA":      {"h1_align": "center", "h1_bold": True,  "h1_italic": False,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": False},
    "Turabian": {"h1_align": "center", "h1_bold": True,  "h1_italic": False,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": False},
    "Harvard":  {"h1_align": "left",   "h1_bold": True,  "h1_italic": False,
                 "h2_align": "left",   "h2_bold": True,  "h2_italic": False},
}

_BIB_HEADINGS: Dict[str, str] = {
    "SBL": "Bibliography", "Chicago": "Bibliography", "MLA": "Works Cited",
    "APA": "References",   "ASA": "References",       "Turabian": "Bibliography",
    "Harvard": "References",
}

_BIB_EXAMPLE: Dict[str, str] = {
    "SBL":      "Author Last, First. Title of Work. Place: Publisher, Year.",
    "Chicago":  "Author Last, First. Title of Work. Place: Publisher, Year.",
    "Turabian": "Author Last, First. Title of Work. Place: Publisher, Year.",
    "MLA":      'Author Last, First. "Title of Article." Journal Name vol. no. (Year): pp–pp.',
    "APA":      "Author, A. A. (Year). Title of work. Publisher. https://doi.org/xxxxx",
    "ASA":      'Author, First. Year. "Title of Article." Journal Name Vol(No):pp–pp.',
    "Harvard":  "Author, A. A. (Year) Title of work. Place: Publisher.",
}


def generate(state: dict) -> bytes:
    """Return a .docx file as bytes. Raises RuntimeError if python-docx is not installed."""
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed.\n\nInstall it with:  pip install python-docx"
        )

    doc = Document()
    cit     = state.get("cit_style", "APA")
    font    = _FONT_MAP.get(state.get("font_pkg", "times"), "Times New Roman")
    size    = _parse_size(state.get("font_size", "12"))
    ls      = state.get("linespace", "2")
    paper   = state.get("paper", "letter")
    margin  = _parse_float(state.get("margin", "1.0"), 1.0)
    title   = state.get("title", "") or "Untitled Essay"
    authors = state.get("authors", "")
    date    = state.get("date", "")
    affil   = state.get("affiliation", "")
    course  = state.get("course", "")
    has_abstract = state.get("has_abstract", False)
    print_bib    = state.get("print_bib", True)
    bib_heading  = state.get("bib_heading", "") or _BIB_HEADINGS.get(cit, "References")
    hs = _HEAD_DEFS.get(cit, _HEAD_DEFS["APA"])

    # ── Page geometry ──────────────────────────────────────────────────────────
    sec = doc.sections[0]
    if paper == "a4":
        sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    else:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    side = Inches(1.25) if cit == "ASA" else Inches(margin)
    sec.top_margin    = Inches(margin)
    sec.bottom_margin = Inches(margin)
    sec.left_margin   = side
    sec.right_margin  = side

    # ── Base styles ────────────────────────────────────────────────────────────
    ls_rule = _ls_rule(ls)
    _configure_normal(doc, font, size, ls_rule)
    _configure_headings(doc, font, size, hs)

    # ── Content ────────────────────────────────────────────────────────────────
    if cit == "MLA":
        _mla_header(doc, authors, affil, course, date, font, size, ls_rule)
        _centered_title(doc, title, font, size, ls_rule)
    else:
        _title_block(doc, title, authors, affil, course, date, font, size, ls_rule)

    if has_abstract:
        _abstract_section(doc, cit, font, size, ls_rule)

    _body_placeholder(doc, font, size, ls_rule)

    if print_bib:
        _bibliography_section(doc, bib_heading, cit, font, size, ls_rule)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_to_odt(docx_bytes: bytes) -> bytes | None:
    """Convert DOCX bytes to ODT bytes via LibreOffice headless. Returns None if unavailable."""
    import os
    import shutil
    import subprocess
    import tempfile

    exe = shutil.which("libreoffice") or shutil.which("soffice")
    if not exe:
        return None
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "input.docx")
        with open(src, "wb") as f:
            f.write(docx_bytes)
        try:
            subprocess.run(
                [exe, "--headless", "--convert-to", "odt", "--outdir", tmpdir, src],
                capture_output=True, timeout=60,
            )
        except (subprocess.TimeoutExpired, FileNotFoundError):
            return None
        out = os.path.join(tmpdir, "input.odt")
        if os.path.exists(out):
            with open(out, "rb") as f:
                return f.read()
    return None


# ── Internal helpers ───────────────────────────────────────────────────────────

def _parse_size(s) -> int:
    try:
        return int(float(str(s).replace("pt", "")))
    except (ValueError, TypeError):
        return 12


def _parse_float(s, default: float) -> float:
    try:
        return float(s)
    except (ValueError, TypeError):
        return default


def _ls_rule(ls: str):
    if ls == "2":
        return WD_LINE_SPACING.DOUBLE
    if ls == "1.5":
        return WD_LINE_SPACING.ONE_POINT_FIVE
    return WD_LINE_SPACING.SINGLE


def _align(s: str):
    return WD_ALIGN_PARAGRAPH.CENTER if s == "center" else WD_ALIGN_PARAGRAPH.LEFT


def _configure_normal(doc, font: str, size: int, ls_rule) -> None:
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(size)
    pf = style.paragraph_format
    pf.line_spacing_rule = ls_rule
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)


def _configure_headings(doc, font: str, size: int, hs: dict) -> None:
    for lvl in (1, 2, 3):
        k = f"h{min(lvl, 2)}"
        try:
            style = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        style.font.name  = font
        style.font.size  = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold   = hs.get(f"{k}_bold", True)
        style.font.italic = hs.get(f"{k}_italic", False)
        style.paragraph_format.alignment    = _align(hs.get(f"{k}_align", "left"))
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after  = Pt(6)


def _add_para(doc, text: str, font: str, size: int, *,
              align: str = "left", ls_rule=None, bold: bool = False, italic: bool = False,
              space_before: float = 0, space_after: float = 0,
              first_line_inches: float = 0.0) -> None:
    if ls_rule is None:
        ls_rule = WD_LINE_SPACING.DOUBLE
    p = doc.add_paragraph()
    p.alignment = _align(align)
    pf = p.paragraph_format
    pf.line_spacing_rule = ls_rule
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.first_line_indent = Inches(first_line_inches)
    run = p.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic


def _title_block(doc, title, authors, affil, course, date, font, size, ls_rule) -> None:
    _add_para(doc, title, font, size, align="center", ls_rule=ls_rule,
              bold=True, space_before=72)
    for text in filter(None, [authors, affil, course, date]):
        _add_para(doc, text, font, size, align="center", ls_rule=ls_rule)
    doc.add_page_break()


def _mla_header(doc, authors, affil, course, date, font, size, ls_rule) -> None:
    for text in filter(None, [authors, affil, course, date]):
        _add_para(doc, text, font, size, ls_rule=ls_rule)


def _centered_title(doc, title, font, size, ls_rule) -> None:
    _add_para(doc, title, font, size, align="center", ls_rule=ls_rule)


def _abstract_section(doc, cit, font, size, ls_rule) -> None:
    if cit in ("APA", "ASA"):
        _add_para(doc, "Abstract", font, size, align="center",
                  ls_rule=ls_rule, bold=True)
        _add_para(doc, "[Write your abstract here. 150–250 words.]",
                  font, size, ls_rule=ls_rule)
        doc.add_page_break()
    else:
        doc.add_heading("Abstract", level=2)
        _add_para(doc, "[Write your abstract here.]", font, size, ls_rule=ls_rule)


def _body_placeholder(doc, font, size, ls_rule) -> None:
    _add_para(
        doc,
        "[Begin your essay here. Replace this paragraph with your introduction.]",
        font, size, ls_rule=ls_rule, first_line_inches=0.5,
    )


def _bibliography_section(doc, heading: str, cit: str, font, size, ls_rule) -> None:
    doc.add_page_break()
    h = doc.add_heading(heading, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if cit in ("SBL", "Chicago", "Turabian") \
                  else WD_ALIGN_PARAGRAPH.LEFT

    entry = doc.add_paragraph()
    pf = entry.paragraph_format
    pf.line_spacing_rule = ls_rule
    pf.left_indent       = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    run = entry.add_run(_BIB_EXAMPLE.get(cit, "Author. Title. Publisher, Year."))
    run.font.name = font
    run.font.size = Pt(size)

    tip = doc.add_paragraph()
    tip.paragraph_format.space_before = Pt(12)
    r = tip.add_run(
        "Tip: Use the Zotero Word plug-in to insert citations and auto-generate this bibliography."
    )
    r.font.name      = font
    r.font.size      = Pt(max(size - 2, 8))
    r.italic         = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
